from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "smart_ims_army_warehouse_quote.docx"


BLUE = "1F4D78"
ACCENT = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C2CC"
DARK = "1F2937"
MUTED = "5B6470"


def money(value):
    return f"INR {value:,.0f}"


def taxable_from_gross(value):
    return int(round(value / 1.18))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{m}"))
        if element is None:
            element = OxmlElement(f"w:{m}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_keep_with_next(paragraph, keep=True):
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:keepNext"))
    if keep and existing is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not keep and existing is not None:
        p_pr.remove(existing)


def add_text(paragraph, text, bold=False, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_keep_with_next(p)
    return p


def add_note_box(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [9360], indent_dxa=120)
    set_cell_margins(table, top=120, bottom=120, start=180, end=180)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_borders(cell, color="CAD6E2", size="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, title, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(body)


def add_amount_summary(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [5600, 3760], indent_dxa=120)
    set_cell_margins(table, top=100, bottom=100, start=160, end=160)
    rows = [
        ("Taxable project value", money(550847)),
        ("GST @ 18%", money(99153)),
        ("Total quoted amount", money(650000)),
        ("Hardware/material/fabrication", "Quoted separately"),
    ]
    for idx, row in enumerate(rows):
        for col in range(2):
            cell = table.cell(idx, col)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if idx == 2:
                set_cell_shading(cell, BLUE)
            elif idx == 3:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if col == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(row[col])
            if idx == 2:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif col == 0:
                run.bold = True
    return table


def add_pricing_table(doc):
    items = [
        ("1", "Requirement study, IMS workflow planning, documentation", 30000),
        ("2", "Core IMS software: login, users, products, locations, inventory database", 130000),
        ("3", "Pick/Put workflows, LED-guided tasks, final review, audit trail, corrections", 95000),
        ("4", "Admin/configuration: users, registration keys, backups, recovery, system health", 45000),
        ("5", "Reporting refinement and print-ready reports", 40000),
        ("6", "RS485 + LED hardware integration software", 110000),
        ("7", "ESP32 firmware integration, controller setup, mapping, test/locate/ping", 75000),
        ("8", "Raspberry Pi packaging, local kiosk/app deployment", 35000),
        ("9", "On-site commissioning, configuration, testing, training", 60000),
        ("10", "Initial support/warranty buffer", 30000),
    ]
    taxable_values = [taxable_from_gross(item[2]) for item in items]
    taxable_values[-1] += 550847 - sum(taxable_values)

    table = doc.add_table(rows=1 + len(items) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [650, 5110, 1800, 1800], indent_dxa=120)
    set_cell_margins(table)

    headers = ("No.", "Component", "Amount excl. GST", "Amount incl. GST")
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True

    for row_idx, item in enumerate(items, start=1):
        no, component, amount = item
        values = (no, component, money(taxable_values[row_idx - 1]), money(amount))
        for col_idx, value in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_idx in (2, 3) else (
                WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            )
            p.add_run(value)

    total_row = 1 + len(items)
    for col_idx in range(4):
        cell = table.cell(total_row, col_idx)
        set_cell_shading(cell, BLUE)
        set_cell_borders(cell, color=BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_idx in (2, 3) else WD_ALIGN_PARAGRAPH.LEFT
        if col_idx == 1:
            text = "Total"
        elif col_idx == 2:
            text = money(550847)
        elif col_idx == 3:
            text = money(650000)
        else:
            text = ""
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)


def add_payment_table(doc):
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1250, 5510, 2600], indent_dxa=120)
    set_cell_margins(table)
    rows = [
        ("Milestone", "Description", "Amount incl. GST"),
        ("40%", "Advance to start finalization and procurement coordination", money(260000)),
        ("30%", "After software packaging and reporting readiness", money(195000)),
        ("20%", "After installation and commissioning", money(130000)),
        ("10%", "After 15-day acceptance period", money(65000)),
    ]
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx == 2 else (
                WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = p.add_run(value)
            if r_idx == 0:
                run.bold = True


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.167


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("LytGuide quotation | One warehouse pilot deployment")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def build_doc():
    doc = Document()
    configure_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("LytGuide Deployment Quotation")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    sub = subtitle.add_run("Smart Inventory Management System | One Warehouse LED-Guided Deployment | Bathinda, Punjab")
    sub.font.size = Pt(12)
    sub.font.color.rgb = RGBColor.from_string(MUTED)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, [2500, 6860], indent_dxa=120)
    set_cell_margins(meta, top=80, bottom=80, start=120, end=120)
    for row_idx, row in enumerate([
        ("Project name", "LytGuide"),
        ("Prepared for", "Army warehouse unit, Bathinda, Punjab"),
        ("Quote date", "09 June 2026"),
        ("Deployment model", "Local Raspberry Pi kiosk with RS485 LED module guidance"),
    ]):
        for col_idx, value in enumerate(row):
            cell = meta.cell(row_idx, col_idx)
            set_cell_borders(cell, color="D5DBE1")
            if col_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            if col_idx == 0:
                run.bold = True

    doc.add_paragraph()
    add_note_box(
        doc,
        "Quoted project value",
        "Total implementation price for LytGuide is INR 6,50,000 inclusive of GST for one local warehouse pilot deployment. Hardware material, cell fabrication, wiring, enclosures, controller boxes, Raspberry Pi peripherals, and replacement parts are quoted separately.",
    )

    add_heading(doc, "Commercial Summary", level=1)
    add_amount_summary(doc)

    add_heading(doc, "Scope Included", level=1)
    add_bullets(doc, [
        "LytGuide local Smart IMS application for one warehouse, intended for approximately 45-50 cells and 2 controllers.",
        "Core inventory software covering users, products, storage locations, stock balances, pick/put tasks, audit transactions, corrections, and recommended actions.",
        "LED-guided inventory operation using RS485-connected ESP32 controller modules and 8x8 matrix LED modules.",
        "Reporting refinement with stock, movement, activity, exception, and print-ready reports.",
        "Admin controls for registration keys, user access, backups, restore flow, system health, controller setup, and cell mapping.",
        "Raspberry Pi local deployment, configuration, commissioning, basic operator/admin training, and initial support.",
    ])

    add_heading(doc, "Price Breakup", level=1)
    add_pricing_table(doc)

    add_heading(doc, "Exclusions and Separately Quoted Items", level=1)
    add_bullets(doc, [
        "LED matrix modules per cell, ESP32 controllers, controller boxes, power supplies, buck converters, wiring, connectors, enclosures, and fabrication material.",
        "Raspberry Pi, display, keyboard/mouse, printer, storage media, and other purchased peripherals.",
        "Travel, stay, local transport, and site-specific logistics, if significant.",
        "Barcode/QR scanning, ERP/accounting integration, mobile app, remote access, multi-warehouse synchronization, and physical push-button confirmation.",
        "Source-code ownership or exclusive IP transfer. The quoted price covers deployment and usage for one warehouse only.",
    ])

    doc.add_page_break()
    add_heading(doc, "Support and Warranty Terms", level=1)
    add_bullets(doc, [
        "Includes 30 days of initial software support and stabilization after commissioning.",
        "Bug fixes within agreed scope are included during the 30-day support period.",
        "New features, report format changes beyond agreed scope, and integrations are change requests.",
        "Replacement parts and failures caused by physical damage, power issues, wiring damage, water/dust ingress, or misuse are charged separately.",
        "Annual maintenance after the initial support period can be quoted separately.",
    ])

    add_heading(doc, "Assumptions", level=1)
    add_bullets(doc, [
        "System is for local warehouse use only and does not require internet/cloud operation.",
        "The first deployment is a pilot/first-unit implementation for one Army warehouse.",
        "The final hardware material and fabrication estimate will be prepared after confirming final cell count, controller count, mounting method, and wiring route.",
        "GST treatment and invoice format should be confirmed with a CA before final billing.",
    ])

    add_heading(doc, "Acceptance", level=1)
    p = doc.add_paragraph()
    p.add_run("The project will be considered commissioned when the local app is installed, configured for the agreed cells/controllers, LED guidance is tested for mapped modules, reports are available, and basic training is completed.")

    doc.add_paragraph()
    sign = doc.add_table(rows=2, cols=2)
    sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(sign, [4680, 4680], indent_dxa=120)
    set_cell_margins(sign, top=120, bottom=120, start=120, end=120)
    labels = [
        ("Prepared by", "Client acknowledgement"),
        ("Name / Signature / Date", "Name / Signature / Date"),
    ]
    for r_idx, row in enumerate(labels):
        for c_idx, value in enumerate(row):
            cell = sign.cell(r_idx, c_idx)
            set_cell_borders(cell, color="D5DBE1")
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(18 if r_idx == 1 else 0)
            run = p.add_run(value)
            if r_idx == 0:
                run.bold = True

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
